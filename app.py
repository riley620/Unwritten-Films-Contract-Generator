from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, json

app = Flask(__name__, static_folder='static', static_url_path='')

GOLD  = RGBColor(0xCE, 0xA0, 0x63)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
MID   = RGBColor(0x59, 0x59, 0x59)
DARK  = RGBColor(0x2C, 0x2C, 0x2C)
LGOLD = RGBColor(0xAA, 0xAA, 0xAA)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '2')
        b.set(qn('w:color'), 'D4B896')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLACK
    run.font.name = 'Arial'
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CEA063')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    run.font.name = 'Arial'

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = MID
    run.font.name = 'Arial'

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = MID
    run.font.name = 'Arial'

def add_info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11.5)
    for i, (label, value) in enumerate(rows):
        row = table.add_row()
        sl = 'F5E6D0' if i % 2 == 0 else 'EDD9B8'
        sv = 'FDFAF6' if i % 2 == 0 else 'FAF5EE'
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = ''
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = DARK; r0.font.name = 'Arial'
        c1.text = ''
        r1 = c1.paragraphs[0].add_run(value or '—')
        r1.font.size = Pt(11); r1.font.color.rgb = MID; r1.font.name = 'Arial'
        for cell, color in [(c0, sl), (c1, sv)]:
            set_cell_bg(cell, color)
            set_cell_borders(cell)

def make_contract(data):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Title
    for text, color in [('VIDEO PRODUCTION', BLACK), ('SERVICES AGREEMENT', GOLD)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(24); r.font.color.rgb = color; r.font.name = 'Arial'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    rs = sub.add_run(f"{data['project_type']} — {data['client_name']}")
    rs.font.size = Pt(11); rs.font.color.rgb = MID; rs.font.name = 'Arial'

    # 1. Parties
    add_h1(doc, '1. PARTIES')
    party_rows = [('Client', f"{data['client_name']} (\"the Client\")")]
    for label, val in data.get('client_rows', []):
        if val and val not in ('—', '--', ''):
            party_rows.append((label, val))
    party_rows += [
        ('Service Provider', f"{data['provider_name']} (\"the Videographer\")"),
        ('ABN',   data.get('provider_abn', '—')),
        ('Email', data.get('provider_email', '—')),
        ('Phone', data.get('provider_phone', '—')),
    ]
    add_info_table(doc, party_rows)
    doc.add_paragraph()

    # 2. Project Details
    add_h1(doc, '2. PROJECT DETAILS')
    add_info_table(doc, [
        ('Project Name',     data.get('project_name', '—')),
        ('Project Type',     data.get('project_type', '—')),
        ('Filming Date(s)',  data.get('filming_dates', '—')),
        ('Filming Location', data.get('filming_location', '—')),
        ('Delivery Date',    data.get('delivery_date', '—')),
    ])
    doc.add_paragraph()

    # 3. Scope
    add_h1(doc, '3. SCOPE OF SERVICES')
    add_body(doc, data.get('scope_desc', ''))
    add_h2(doc, 'Inclusions')
    add_bullet(doc, f"Up to {data.get('revisions', '2')} round(s) of revisions based on Client feedback")
    add_bullet(doc, f"Raw footage: {data.get('raw_footage', 'Not included')}")
    add_bullet(doc, 'Professional filming equipment, audio and lighting as required')
    add_bullet(doc, 'Colour grade, audio mix and titles/graphics as agreed')
    doc.add_paragraph()

    # 4. Fees
    add_h1(doc, '4. FEES & PAYMENT')
    fee_table = doc.add_table(rows=1, cols=3)
    fee_table.style = 'Table Grid'
    fee_table.autofit = False
    fee_table.columns[0].width = Cm(9)
    fee_table.columns[1].width = Cm(4.5)
    fee_table.columns[2].width = Cm(3.5)
    hdr = fee_table.rows[0]
    for j, txt in enumerate(['Description', 'Amount (AUD)', 'GST Incl.']):
        c = hdr.cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GOLD; r.font.name = 'Arial'
        set_cell_bg(c, '1A1A1A')
        set_cell_borders(c)

    total = 0
    for i, fee in enumerate(data.get('fee_rows', [])):
        if not fee.get('desc'):
            continue
        row = fee_table.add_row()
        sh = 'FDFAF6' if i % 2 == 0 else 'FAF5EE'
        try:
            amt = float(str(fee.get('amt', '0')).replace('$', '').replace(',', ''))
            total += amt
            amt_str = f"${amt:,.2f}"
        except:
            amt_str = fee.get('amt') or '—'
        for j, txt in enumerate([fee['desc'], amt_str, 'Yes']):
            c = row.cells[j]
            c.text = ''
            r = c.paragraphs[0].add_run(txt)
            r.font.size = Pt(11); r.font.color.rgb = MID; r.font.name = 'Arial'
            set_cell_bg(c, sh)
            set_cell_borders(c)

    tot_row = fee_table.add_row()
    for j, txt in enumerate(['TOTAL', f"${total:,.2f}" if total > 0 else 'TBC', '']):
        c = tot_row.cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK; r.font.name = 'Arial'
        set_cell_bg(c, 'F5E6D0')
        set_cell_borders(c)

    doc.add_paragraph()
    add_h2(doc, 'Payment Terms')
    add_bullet(doc, f"Deposit of {data.get('deposit_pct', '50')}% due upon signing of this Agreement")
    add_bullet(doc, f"Balance due within {data.get('balance_days', '14')} days of final delivery")
    add_bullet(doc, f"Payment via {data.get('payment_method', 'EFT')}")
    add_bullet(doc, 'Late payments may attract interest of 1.5% per month')
    doc.add_paragraph()

    # 5. Cancellation
    add_h1(doc, '5. CANCELLATION & VARIATION')
    add_bullet(doc, f"Cancellations with more than {data.get('cancel_notice', '14')} days notice: full deposit refunded")
    add_bullet(doc, f"Cancellations with less than {data.get('cancel_notice', '14')} days notice: deposit is {data.get('cancel_deposit', 'non-refundable')}")
    add_bullet(doc, 'Material changes to scope must be agreed in writing and may attract additional fees')
    doc.add_paragraph()

    # 6. Terms & Conditions
    clauses = data.get('clauses', [])
    if clauses:
        add_h1(doc, '6. TERMS & CONDITIONS')
        for clause in clauses:
            add_h2(doc, clause['title'])
            for line in clause['lines']:
                if line.startswith('•'):
                    add_bullet(doc, line[1:].strip())
                else:
                    add_body(doc, line)
        doc.add_paragraph()

    # 7. General
    add_h1(doc, '7. GENERAL')
    add_bullet(doc, 'This Agreement is governed by the laws of Queensland, Australia')
    add_bullet(doc, 'Disputes will first be attempted to be resolved by negotiation in good faith')
    add_bullet(doc, 'This document constitutes the entire agreement and supersedes all prior discussions')
    add_bullet(doc, 'Amendments must be in writing and signed by both parties')
    doc.add_paragraph()

    # 8. Signatures
    add_h1(doc, '8. SIGNATURES')
    add_body(doc, 'By signing below, both parties agree to be bound by the terms of this Agreement.')
    doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(8.5)
    sig_table.columns[1].width = Cm(8.5)
    for j, lbl in enumerate([f"FOR {data['client_name'].upper()}", data['provider_name'].upper()]):
        cell = sig_table.rows[0].cells[j]
        cell.text = ''
        set_cell_bg(cell, 'FDFAF6')
        set_cell_borders(cell)
        cell.add_paragraph()
        for line in ['Signature: ___________________________', 'Name: ________________________________',
                     'Title: _______________________________', 'Date: ________________________________']:
            p = cell.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11); r.font.color.rgb = MID; r.font.name = 'Arial'
            p.paragraph_format.space_before = Pt(4)
        p_lbl = cell.add_paragraph()
        r_lbl = p_lbl.add_run(lbl)
        r_lbl.bold = True; r_lbl.font.size = Pt(10); r_lbl.font.color.rgb = GOLD; r_lbl.font.name = 'Arial'
        p_lbl.paragraph_format.space_before = Pt(8)

    doc.add_paragraph()
    disc = doc.add_paragraph('This Agreement was prepared as a general template. Both parties should seek independent legal advice before signing.')
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc.runs[0].italic = True
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].font.color.rgb = LGOLD

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# Clause content lookup
CLAUSE_CONTENT = {
    'ip-ownership': {
        'title': 'Intellectual Property & Copyright',
        'lines': [
            'Upon receipt of full payment, the Client will own the copyright in the final delivered video(s). The Videographer retains copyright in all raw footage until full payment is received.',
        ]
    },
    'portfolio-use': {
        'title': 'Portfolio & Promotional Use',
        'lines': [
            'The Videographer may use excerpts or still frames from the final delivered video(s) for portfolio and promotional purposes, unless the Client objects in writing within 14 days of delivery. Any use involving identifiable minors requires separate written consent.',
        ]
    },
    'confidentiality': {
        'title': 'Confidentiality',
        'lines': [
            'Both parties agree to keep confidential any proprietary or sensitive information shared during this engagement and not to disclose it to third parties without prior written consent, except as required by law.',
        ]
    },
    'liability': {
        'title': 'Liability & Insurance',
        'lines': [
            '• The Videographer holds public liability insurance. Certificate of Currency available on request.',
            "• The Videographer's liability is limited to the total fees paid under this Agreement.",
            '• The Client is responsible for ensuring a safe working environment at all filming locations.',
        ]
    },
    'privacy': {
        'title': 'Privacy',
        'lines': [
            'All personal information is handled in accordance with the Privacy Act 1988 (Cth). Raw footage will be securely deleted within 90 days of final delivery unless otherwise agreed in writing.',
        ]
    },
    'child-safety': {
        'title': 'Child Safety',
        'lines': [
            '• The Videographer holds a current Working with Children Check (Blue Card) under the Working with Children (Risk Management and Screening) Act 2000 (Qld).',
            '• A responsible staff member will be present at all times during filming involving minors.',
            '• The Videographer will not be left unsupervised with minors at any time.',
        ]
    },
    'consent-school': {
        'title': 'Student Consent (Managed by School)',
        'lines': [
            'The School is solely responsible for obtaining and managing all parental/guardian consent forms prior to filming, in compliance with the Education (General Provisions) Act 2006 (Qld). Consent records will be retained for a minimum of 7 years.',
            'The School will notify the Videographer prior to filming of any students who do not have consent to be filmed.',
        ]
    },
    'backup-policy': {
        'title': 'Footage Backup Policy',
        'lines': [
            'The Videographer will maintain at least two backup copies of all raw footage until final delivery and client sign-off.',
        ]
    },
    'weather-clause': {
        'title': 'Weather & Force Majeure',
        'lines': [
            "If filming is postponed due to adverse weather, natural disaster, government direction, or other circumstances beyond either party's reasonable control, both parties will agree a new filming date at no additional charge.",
        ]
    },
    'site-safety': {
        'title': 'Mine Site Safety',
        'lines': [
            '• The Client must arrange all necessary site induction, access permits and PPE requirements prior to the filming date.',
            '• The Videographer will comply with all site safety, HSE, and access requirements as directed by the Client.',
        ]
    },
    'likeness-rights': {
        'title': 'Likeness & Publicity Rights',
        'lines': [
            "The Client grants permission for the Videographer to film and reproduce the athlete's likeness for the purposes of this project only. Any commercial use beyond the agreed deliverables requires a separate written agreement.",
        ]
    },
    'access-requirements': {
        'title': 'Event Access Requirements',
        'lines': [
            '• The Client will provide the Videographer with appropriate media/crew accreditation and access to all areas required for filming as agreed.',
            '• Any restrictions on access or filming areas must be communicated in writing prior to the event date.',
        ]
    },
}

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()

        # Build clauses list from selected keys
        selected_clauses = data.get('clauses', {})
        clauses = [CLAUSE_CONTENT[k] for k in selected_clauses if selected_clauses[k] and k in CLAUSE_CONTENT]
        data['clauses'] = clauses

        buf = make_contract(data)

        client_name = data.get('client_name', 'Client').replace(' ', '_')
        project_type = data.get('project_type', 'Contract').replace(' ', '_')
        filename = f"Unwritten_Films_{project_type}_{client_name}.docx"

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
